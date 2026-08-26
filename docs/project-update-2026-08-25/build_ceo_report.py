from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "presentation" / "scratch" / "assets" / "vitarich-ceo-update"
OUTPUT_DIR = ROOT / "output"
OUTPUT_PATH = OUTPUT_DIR / "Vita-FMS-Total-Project-Update-May-August-2026.docx"

GREEN = "144D32"
GREEN_2 = "2F6B4F"
GOLD = "9B7B1B"
GOLD_LIGHT = "E8DCA8"
INK = "1F2933"
MUTED = "5B6670"
LIGHT = "F2F6F3"
LIGHT_GOLD = "F7F3E5"
WHITE = "FFFFFF"
RED = "A13D3D"
AMBER = "A76516"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Segoe UI", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, *, size=10.5, bold=False, color=INK, italic=False, font="Segoe UI"):
    run = paragraph.add_run(text)
    set_font(run, font, size, bold, color, italic)
    return run


def add_heading(doc, text, level=1, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_text(p, kicker.upper(), size=8.5, bold=True, color=GOLD, font="Bahnschrift")
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, "Bahnschrift", 20 if level == 1 else 14, True, GREEN if level == 1 else INK)
    return p


def add_body(doc, text, *, bold_lead=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text, level=0, color=INK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.18 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    add_text(p, text, size=10, color=color)
    return p


def add_callout(doc, title, body, fill=LIGHT, accent=GREEN):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_cell_width(table.cell(0, 0), 0.10)
    set_cell_width(table.cell(0, 1), 6.25)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    set_cell_margins(table.cell(0, 0), 80, 20, 80, 20)
    set_cell_margins(table.cell(0, 1), 150, 180, 150, 180)
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, size=11, bold=True, color=accent, font="Bahnschrift")
    p2 = table.cell(0, 1).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    add_text(p2, body, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_kpi_row(doc):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    values = [
        ("98", "integrated commits\nin dev-main"),
        ("469", "files changed\nsince May 1"),
        ("4", "monthly delivery\ncycles covered"),
        ("5", "portfolio categories\nmapped end to end"),
    ]
    for idx, (value, label) in enumerate(values):
        cell = table.cell(0, idx)
        set_cell_width(cell, 1.59)
        set_cell_shading(cell, LIGHT if idx % 2 == 0 else LIGHT_GOLD)
        set_cell_margins(cell, 150, 120, 150, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_text(p, value, size=25, bold=True, color=GREEN, font="Bahnschrift")
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_text(p2, label, size=8.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_module_table(doc, rows, widths=(1.35, 1.75, 2.45, 0.85)):
    headers = ("Module", "What changed", "Connected to / business role", "Status")
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, GREEN)
        set_cell_margins(cell, 110, 100, 110, 100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 3 else WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, header, size=8.5, bold=True, color=WHITE, font="Bahnschrift")
    for row_index, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx], 105, 100, 105, 100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cells[idx], "F8FAF8")
            if idx == 3:
                status_color = GREEN if value == "Integrated" else AMBER if "Progress" in value or "Branch" in value else RED
                p = cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(p, value, size=8.3, bold=True, color=status_color)
            else:
                p = cells[idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                add_text(p, value, size=8.6, bold=(idx == 0), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_flow(doc, nodes, caption=None):
    table = doc.add_table(rows=1, cols=len(nodes) * 2 - 1)
    table.autofit = False
    for i, node in enumerate(nodes):
        cell = table.cell(0, i * 2)
        set_cell_width(cell, 0.94)
        set_cell_shading(cell, LIGHT if i % 2 == 0 else LIGHT_GOLD)
        set_cell_margins(cell, 125, 75, 125, 75)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, node, size=8.2, bold=True, color=GREEN, font="Bahnschrift")
        if i < len(nodes) - 1:
            arrow = table.cell(0, i * 2 + 1)
            set_cell_width(arrow, 0.15)
            p2 = arrow.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p2, "→", size=12, bold=True, color=GOLD)
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        add_text(p, caption, size=8.2, italic=True, color=MUTED)
    return table


def add_timeline_table(doc):
    rows = [
        ("May", "24", "Process visibility and control", "Hatchery views/APIs, receiving trace, system-adoption reporting, permission templates, and standardized operational tables."),
        ("June", "27", "Transaction and inventory foundations", "Trace reversal/validation, Breeder growing and grading, farm-aware filters, Item Group/UoM masters, batching, Goods Receipt, and Goods Issue."),
        ("July", "26", "Connected Broiler operating model", "DOC Placement, Flock Card, Farm Setup Wizard, approvals, Harvest & Delivery, Clean up, Inventory Transfer, warehouse reporting, and posting lineage."),
        ("August", "21", "Lifecycle governance and reliability", "Cycle Master, farm-cycle closure, FMS-scoped access, notifications, network and theme resilience, delivery refinements, and item hierarchy work."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    headers = ("Month", "Commits", "Executive theme", "Representative outcomes")
    widths = (0.65, 0.60, 1.65, 3.40)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, GREEN)
        set_cell_margins(cell, 100, 90, 100, 90)
        add_text(cell.paragraphs[0], header, size=8.5, bold=True, color=WHITE, font="Bahnschrift")
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i], 110, 90, 110, 90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cells[i], "F8FAF8")
            p = cells[i].paragraphs[0]
            if i in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value, size=8.6, bold=(i in (0, 2)), color=GREEN if i == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_legend(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "Status legend: ", size=9, bold=True, color=INK)
    add_text(p, "Integrated", size=9, bold=True, color=GREEN)
    add_text(p, " = committed to dev-main; ", size=9, color=MUTED)
    add_text(p, "Branch-only", size=9, bold=True, color=AMBER)
    add_text(p, " = committed but not integrated; ", size=9, color=MUTED)
    add_text(p, "In Progress", size=9, bold=True, color=AMBER)
    add_text(p, " = local uncommitted work; ", size=9, color=MUTED)
    add_text(p, "Verification pending", size=9, bold=True, color=RED)
    add_text(p, " = code exists but live deployment or runtime evidence is incomplete.", size=9, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Segoe UI")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Segoe UI")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    for level, size in ((1, 20), (2, 14), (3, 11)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Bahnschrift"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Bahnschrift")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Bahnschrift")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN if level == 1 else INK)

    return doc


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p, "VITA FMS  •  CEO PROJECT UPDATE", size=7.5, bold=True, color=GREEN, font="Bahnschrift")
        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Inches(6.8))
        table.autofit = False
        p1 = table.cell(0, 0).paragraphs[0]
        add_text(p1, "Internal management document", size=7.5, color=MUTED)
        p2 = table.cell(0, 1).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p2, "Reporting period: 01 May–25 August 2026", size=7.5, color=MUTED)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    p.add_run().add_picture(str(ASSET_DIR / "vitarich-logo.png"), width=Inches(0.78))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_text(p, "VITA FMS", size=11, bold=True, color=GOLD, font="Bahnschrift")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "Total Project Update", size=34, bold=True, color=GREEN, font="Bahnschrift")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    add_text(p, "From process digitization to a connected farm-management operating platform", size=15, color=INK, font="Georgia")
    p = doc.add_paragraph()
    p.add_run().add_picture(str(ASSET_DIR / "integrated-campus.png"), width=Inches(6.82))
    p.paragraph_format.space_after = Pt(12)
    meta = doc.add_table(rows=2, cols=2)
    meta.autofit = False
    labels = (("Reporting period", "01 May–25 August 2026"), ("Prepared for", "Chief Executive Officer"))
    for r, (label, value) in enumerate(labels):
        set_cell_width(meta.cell(r, 0), 1.35)
        set_cell_width(meta.cell(r, 1), 5.0)
        set_cell_margins(meta.cell(r, 0), 80, 60, 80, 60)
        set_cell_margins(meta.cell(r, 1), 80, 60, 80, 60)
        add_text(meta.cell(r, 0).paragraphs[0], label.upper(), size=7.5, bold=True, color=GOLD, font="Bahnschrift")
        add_text(meta.cell(r, 1).paragraphs[0], value, size=9.5, bold=True, color=INK)
    doc.add_page_break()

    add_heading(doc, "Executive summary", kicker="CEO readout")
    add_callout(
        doc,
        "Headline",
        "Between May and August, Vita FMS moved from a collection of operational screens toward an increasingly connected platform spanning Hatchery, Breeder, Broiler, Inventory, and shared governance. The strongest progress is the connection of operational events to farm, warehouse, batch, cycle, permission, approval, reporting, and notification structures.",
    )
    add_kpi_row(doc)
    add_body(doc, "The most material business advance is the Broiler lifecycle: DOC Placement can establish farm and building cycles; Growing & Farm Condition carries daily flock and inventory activity; Harvest & Delivery uses eligible flock and batch data; Clean up closes participating cycles; and inventory reports reconcile the resulting ledger movements.")
    add_body(doc, "The second major advance is the Inventory backbone. Item, group, UoM, warehouse, batch, receipt, issue, transfer, audit, and warehouse-report capabilities now provide a common transaction layer used directly by operating modules.")
    add_body(doc, "Governance also matured through FMS-scoped user roles, standalone permission management, approval workflows, centralized notification rules, safer navigation enforcement, farm-authorized data filtering, and shared user-experience improvements.")
    add_callout(
        doc,
        "Executive caution",
        "Integrated code is not the same as a production-complete capability. Several migrations, Supabase/RPC changes, browser flows, email delivery, and branch/local work still require controlled deployment and live verification. This report therefore presents delivery status explicitly and avoids treating unverified work as released.",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "Scope and evidence", kicker="Method")
    add_bullet(doc, "GitHub remotes were refreshed on 25 August 2026; the integrated baseline is origin/dev-main.")
    add_bullet(doc, "The integrated period contains 98 commits: May 24, June 27, July 26, and August 21.")
    add_bullet(doc, "The net change from the 29 April baseline to 25 August covers 469 files, 85,930 insertions, and 7,483 deletions. These are engineering-volume indicators, not business-value KPIs.")
    add_bullet(doc, "Branch-only work includes two Breeder commits on dev-baja and a Week Lock module on dev-cris.")
    add_bullet(doc, "The initial worktree snapshot contained 45 modified tracked files and seven pre-existing untracked feature paths; these are reported as In Progress.")
    add_bullet(doc, "Repository history, routes, settings, data-access code, SQL, and prior validation notes were reconciled to identify what each module connects to and where runtime verification remains incomplete.")
    add_status_legend(doc)

    add_heading(doc, "Portfolio architecture", kicker="What connects to what")
    add_flow(
        doc,
        ["Farm & user context", "FMS operating modules", "Inventory ledger", "Reports & trace", "Notifications & approvals"],
        "Shared master data and governance frame every business workflow; operational posting then flows into ledger, reporting, trace, and communication controls.",
    )
    portfolio_rows = [
        ("Shared Platform", "Farm setup, users, permissions, approvals, notifications, navigation, reliability", "Controls who can operate, which farm/FMS is in scope, and how actions are governed and communicated.", "Integrated"),
        ("Inventory", "Masters, batches, stock in/out, transfers, audit, warehouse reporting", "Supplies item/warehouse/batch data and receives postings from operating modules.", "Integrated"),
        ("Broiler", "DOC Placement, Cycle Master, Growing & Farm Condition, Harvest & Delivery, Clean up", "Connected lifecycle from placement through daily growing, delivery, clean-out, and cycle closure.", "Integrated"),
        ("Hatchery", "Receiving through DOC Dispatch and Disposal", "Sequential egg-to-chick process with shared trace, inventory, and forthcoming process wizard.", "Integrated"),
        ("Breeder", "Placement, population, laying, health, dispatch, clean-up, reports", "Operational chain with additional transfer/history/cleanup work still branch-only.", "Integrated"),
    ]
    add_module_table(doc, portfolio_rows)
    doc.add_page_break()

    add_heading(doc, "Delivery timeline", kicker="May to August")
    add_timeline_table(doc)
    add_body(doc, "The sequence matters: the project first improved process visibility and access control, then built transaction foundations, then connected Broiler operations, and finally strengthened lifecycle governance and reliability. This is a platform-building pattern rather than a collection of isolated screen changes.")

    add_heading(doc, "Shared platform and governance", kicker="Cross-FMS")
    shared_rows = [
        ("Farm Management & Setup Wizard", "Farm create/edit workflows, warehouse assignment, Building/Pen structure, and default feed/receiving/disposal warehouses.", "Feeds every farm-scoped form, warehouse selection, Broiler cycle, and recipient-routing decision.", "Integrated"),
        ("User Management", "FMS type and three-level user model prepared with server-side management boundaries.", "Connects identity to Broiler, Breeder, or Hatchery scope and to manageable users.", "Integrated"),
        ("User Permissions", "Standalone searchable editor, bulk allow/remove, server-preloaded navigation catalog, and layout-matching loading states.", "NavFolders, sidebar visibility, direct-route guards, and API authorization.", "Integrated"),
        ("Approval", "Approval templates, management, activation, and password-reset authorization improvements.", "Shared approval states and supervisors used by supported business and administrative actions.", "Integrated"),
        ("Notification Setup", "Central event catalog, rules, outbox/delivery processing, inbox, templates, placeholders, retries, and email controls.", "Module Post/Edit/Void events → permission-aware recipients → in-app and optional email delivery.", "Integrated"),
        ("Trace & Validate", "Transaction graph, node arrangement, modal detail, reversal/cancellation visibility, and route validation.", "Hatchery Receiving and downstream process records; disposal and reverse-transaction paths.", "Integrated"),
        ("Navigation, Search & UI", "FMS-scoped navigation, new-document actions, route repairs, breadcrumbs/browser titles, theme-aware controls, and offline banner.", "Applies across all modules and reduces route, permission, and usability inconsistency.", "Integrated"),
        ("System Adoption Report", "Operational usage visualization and farm-aware report filtering.", "Management view across active system usage.", "Integrated"),
        ("Week Lock", "Week Lock administration was developed on dev-cris.", "Would constrain document activity by configured accounting/operational periods.", "Branch-only"),
    ]
    add_module_table(doc, shared_rows)
    add_callout(doc, "Governance connection", "Notifications do not grant access. The central design rechecks module permission, FMS scope, farm assignment, and configured rules before delivery; a rule with no match is designed to be a safe no-op.")

    add_heading(doc, "Inventory backbone", kicker="Shared transaction layer")
    add_flow(doc, ["Item / Group / UoM", "Warehouse & Batch", "Stock In / Out / Transfer", "inventory_postings", "Audit & Warehouse Report"], "Every posted movement is intended to become a traceable signed ledger entry that can be reconciled and linked back to its source document.")
    inventory_rows = [
        ("Item Master & Item Group", "Item CRUD, FMS grouping, one-level subgroup hierarchy, and dependent subgroup selection.", "Supplies item identity to receipt, issue, transfer, Flock Card feed/mortality, and operating modules.", "In Progress"),
        ("UoM Master & Conversion", "New unit and conversion masters connected to item and transaction forms.", "Normalizes transaction quantities before posting.", "Integrated"),
        ("Warehouse Master", "Farm assignment, building/pen relationships, and default warehouse roles.", "Provides authorized sources/destinations for all inventory and farm workflows.", "Integrated"),
        ("Batch Manager", "Batch number/expiry support and canonical batch availability patterns.", "Feeds Goods Receipt, Goods Issue, Transfer, Flock Card, Harvest & Delivery, and Clean up allocations.", "Integrated"),
        ("Item Stock In", "Goods Receipt forms, batching, posting, draft/post workflow, DR reference, and line-level group/subgroup presentation.", "Vendor/farm/warehouse/batch → inventory_postings; direct-post refinements are local WIP.", "In Progress"),
        ("Item Stock Out", "General Goods Issue plus shared forms used by Harvest & Delivery and Clean up.", "Manual GI, BR-DR, and BR-CU document types → dedicated tables and signed OUT postings.", "Integrated"),
        ("Inventory Transfer", "Origin/destination warehouse posting and batch/on-hand validation.", "Creates paired warehouse movements and feeds the common ledger; direct-post refinement is local WIP.", "In Progress"),
        ("Inventory Audit", "Farm, warehouse, date filters, visible posting ID, and descending order.", "Direct view of inventory_postings for transaction review.", "Integrated"),
        ("Warehouse Report", "Beginning/running balance, batch separation, farm-authorized warehouse search, Excel/PDF export, and source links.", "Reconciles receipts, issues, transfers, DOC Placement, Flock Card, Harvest & Delivery, and Clean up postings.", "Integrated"),
    ]
    add_module_table(doc, inventory_rows)
    add_callout(doc, "Business significance", "The inventory ledger is the common financial-operational spine. A quantity shown in an operating module should reconcile with inventory_postings and the Warehouse Report rather than being recalculated independently in each screen.")
    doc.add_page_break()

    add_heading(doc, "Broiler lifecycle", kicker="End-to-end operating chain")
    add_flow(doc, ["DOC Placement", "Farm / Building Cycle", "Growing & Farm Condition", "Harvest & Delivery", "Clean up"], "Cycle Master monitors the lifecycle; settings govern each stage; inventory_postings records stock consequences.")
    broiler_rows = [
        ("DOC Placement Settings", "Farm-scoped item mapping for good, dead-on-arrival, and rejected chicks; excluded-cycle buildings.", "Direct settings module for DOC Placement; controls classification item mapping and farm-cycle exceptions.", "Integrated"),
        ("DOC Placement", "Per-line building placement, date and age rules, actual-received logic, direct posting, source lineage, remarks, and cycle creation.", "Creates or joins Farm/Building cycles, creates Flock Cards, and posts DOC inventory through shared Goods Receipt structures.", "Integrated"),
        ("Cycle Master", "Read-only farm cycle view with participating-building progress and closure state.", "Cycles begin from DOC Placement and close only after participating Flock Cards complete posted Clean up.", "Integrated"),
        ("Growing & Farm Condition Settings", "Farm-specific feed group, automatic feed-batch mode, and other Flock Card configuration.", "Direct settings module for the daily Growing & Farm Condition / Flock Card workflow.", "In Progress"),
        ("Growing & Farm Condition", "Daily Flock Card grid, placement origin, feed/mortality batches, on-hand checks, bird metrics, guidelines, exports, and loading/action feedback.", "Consumes configured items and batches; records feed and mortality/thinning movements; supplies age/body weight to downstream modules.", "In Progress"),
        ("Harvest & Delivery Settings", "Farm-scoped target delivery age and batch auto-selection.", "Direct settings module for Harvest & Delivery eligibility and allocation behavior.", "Integrated"),
        ("Harvest & Delivery", "Eligible-building selection, flock age/body weight, multi-batch allocation, transport fields, receipt, and dedicated posting tables.", "Reads Flock Card state and creates BR-DR inventory issues through the shared Goods Issue engine.", "Integrated"),
        ("Clean up Settings", "Target clean-up age and batch auto-selection by farm.", "Direct settings module for Clean up eligibility and placement-batch selection.", "Integrated"),
        ("Clean up", "Dedicated BR-CU document, quantity/variance posting, receipt, report, and lifecycle close-out.", "Uses remaining flock/batch position, posts clean-out movements, then closes eligible cycle participation.", "Integrated"),
    ]
    add_module_table(doc, broiler_rows)
    add_callout(doc, "Terminology note", "The current business navigation label is DOC Placement. Its implementation route remains /inv/doc-receiving, and its settings route remains /a_dean/doc-receiving-settings. This is separate from Hatchery Receiving at /a_dean/receiving.", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "Hatchery process chain", kicker="Egg to DOC dispatch")
    add_flow(doc, ["Receiving", "Classification / Storage", "Pre-Warm / Setter", "Transfer / Hatcher", "Pullout / DOC Dispatch"], "Trace and inventory controls sit across the chain; the Process Wizard is currently local work in progress.")
    hatchery_rows = [
        ("Receiving", "Manual receiving workflow, farm selection, approval readiness, and trace integration.", "Entry point for hatchery inventory/process lineage; current Shipped To filtering is local WIP.", "In Progress"),
        ("Egg Classification", "Classification features, list/view/API improvements, and standardized table behavior.", "Consumes Receiving context and organizes eggs before storage/process steps.", "Integrated"),
        ("Egg Storage", "New/view/list workflows and inventory-oriented refinements.", "Receives classified eggs and supplies pre-warming.", "Integrated"),
        ("Pre-Warming", "Process forms, views, and route support.", "Bridges stored eggs to Setter processing.", "Integrated"),
        ("Egg Setter", "Setter process forms, APIs, edit/void support, and allocation refinements.", "Feeds Egg Transfer and Hatcher stages.", "Integrated"),
        ("Egg Transfer", "Transfer process forms/APIs and edit/void support.", "Moves set eggs into Hatcher processing.", "Integrated"),
        ("Egg Hatcher", "Hatcher forms/views and process updates.", "Feeds chick pullout.", "Integrated"),
        ("Chick Pullout", "Pullout forms/views and workflow improvements.", "Feeds DOC Classification.", "Integrated"),
        ("DOC Classification", "Classification totals, views/APIs, schema fixes, and notification readiness work.", "Creates graded DOC outcomes used before dispatch/disposal.", "Integrated"),
        ("DOC Dispatch", "List, create, view, post, print, schema/RPC fixes, and notification event integration.", "Final Hatchery dispatch; operationally supplies downstream receiving/placement references.", "Integrated"),
        ("Disposal", "Disposal entry and trace reversal support.", "Captures non-usable process outcomes and inventory consequences.", "Integrated"),
        ("Hatchery Process Wizard", "Guided cross-module workflow is present in local uncommitted routes and shared repositories.", "Intended to reduce repeated selection and carry process context across Hatchery stages.", "In Progress"),
    ]
    add_module_table(doc, hatchery_rows)

    add_heading(doc, "Breeder operating chain", kicker="Placement to close-out")
    add_flow(doc, ["Placement", "Population Record", "Laying Production", "Vaccination / Medication", "Dispatch / Clean-Up"], "The Breeder Card and reports aggregate lifecycle activity; transfer/history improvements remain branch-only.")
    breeder_rows = [
        ("Placement", "Placement forms, source policies, cycle creation, warehouse linkage, and table refinements.", "Starts the Breeder Card lifecycle and supplies population baseline.", "Integrated"),
        ("Population Record", "Growing forms, grading, APIs, filtering, and farm-aware operational records.", "Updates flock population and performance after placement.", "Integrated"),
        ("Laying Production", "Laying forms, APIs, list improvements, and Card integration.", "Adds production results to the Breeder lifecycle.", "Integrated"),
        ("Vaccination & Medication", "Forms, tables, edit support, and database rules for flock health interventions.", "Records health activity against Breeder flocks.", "Integrated"),
        ("Breeder Card / Reports", "Card export, daily performance API, and report screens.", "Consolidates placement, population, production, health, dispatch, and cleanup history.", "Integrated"),
        ("Breeder Logistics", "Dispatch, redesigned Clean-Up, Transfer, detailed History, print, RLS, and Card linkage are present in dev-baja branch work.", "Moves or reduces flock population, improves lifecycle traceability, and supports final close-out.", "Branch-only"),
    ]
    add_module_table(doc, breeder_rows)
    doc.add_page_break()

    add_heading(doc, "Work in progress and release boundary", kicker="As of 25 August")
    add_callout(doc, "Local worktree", "At the start of this review, 45 tracked files were modified and seven pre-existing feature paths were untracked. The largest themes are Flock Card feed/mortality allocation, Item Group/Subgroup persistence, Item Master subgroup use, Goods Receipt/Transfer direct posting, Hatchery Process Wizard, manual Hatchery Receiving farm filtering, shared sidebar/navigation refinements, and notification catalog extensions.", fill=LIGHT_GOLD, accent=GOLD)
    wip_rows = [
        ("Flock Card", "Feed Group requirement, feed-type-to-batch filtering, RPC signature alignment, mandatory feed/mortality allocation validation, grid/export refinements.", "Growing & Farm Condition Settings → Item Group/Subgroup → Batch availability → inventory_postings.", "In Progress"),
        ("Item hierarchy", "One-level subgroup persistence, server-authorized mutation, Item Master subgroup assignment, and notification event support.", "Item Group → Item Master → Goods Receipt display → Flock Card feed filtering.", "In Progress"),
        ("Inventory direct posting", "Goods Receipt and Inventory Transfer can retain permissive drafts while applying strict validation only when posting.", "New form → shared mutation → inventory ledger → audit/report.", "In Progress"),
        ("Hatchery Process Wizard", "New wizard route, API paths, shared data repositories, and form updates across Hatchery steps.", "Receiving → Classification → Storage → Pre-Warm → Setter → Transfer → Hatcher → Pullout → DOC Classification/Dispatch.", "In Progress"),
        ("Breeder lifecycle expansion", "Transfer, history, dispatch print, revised cleanup, and additional card metrics on dev-baja.", "Placement/Card → Transfer/History → Dispatch/Cleanup.", "Branch-only"),
        ("Week Lock", "Administrative period-lock capability on dev-cris.", "Would govern which documents can be changed within a locked week.", "Branch-only"),
    ]
    add_module_table(doc, wip_rows)

    doc.add_page_break()
    add_heading(doc, "Delivery confidence and principal risks", kicker="Management view")
    risks = [
        ("Deployment evidence gap", "Several checked-in SQL/RPC changes are not confirmed as applied to Supabase; repository success should not be reported as production success.", "Create a migration register with environment, owner, execution date, rollback, and smoke-test evidence."),
        ("Branch and worktree concentration", "Material Breeder work is branch-only and current local work spans many modules, increasing integration and regression risk.", "Package work into smaller PRs by FMS/module and merge behind a documented release train."),
        ("Live workflow verification", "Static TypeScript/lint checks are common; browser, login, permissions, concurrency, and ledger reconciliation are not consistently evidenced.", "Adopt a release checklist covering role-based login, farm scope, draft/post/void, batch/on-hand, report reconciliation, and browser QA."),
        ("Notification activation", "Central architecture exists, but email provider authentication failed in prior live testing and some modules still have farm-identity gaps.", "Complete canonical farm-ID audit, deploy outbox SQL, configure provider credentials, and verify no-rule safe no-op plus retry deduplication."),
        ("Change traceability", "Several commits have low-information subjects such as dates, letters, or 'build', making executive and release reconstruction difficult.", "Require PR titles and commit subjects to name FMS, module, action, and business outcome."),
        ("Data-contract drift", "Legacy names and routes can obscure business meaning, especially Hatchery Receiving versus Broiler DOC Placement.", "Maintain a canonical module catalog with business label, route, owner, settings module, tables/RPCs, and notification keys."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = (1.30, 2.45, 2.65)
    for i, h in enumerate(("Risk", "Why it matters", "Recommended control")):
        set_cell_width(table.cell(0, i), widths[i])
        set_cell_shading(table.cell(0, i), GREEN)
        set_cell_margins(table.cell(0, i), 100, 100, 100, 100)
        add_text(table.cell(0, i).paragraphs[0], h, size=8.5, bold=True, color=WHITE, font="Bahnschrift")
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(risks, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i], 105, 100, 105, 100)
            if r_idx % 2 == 0:
                set_cell_shading(cells[i], "F8FAF8")
            add_text(cells[i].paragraphs[0], value, size=8.5, bold=(i == 0), color=INK)

    doc.add_page_break()
    add_heading(doc, "Recommended CEO priorities", kicker="Next 30–60 days")
    priorities = [
        ("1", "Stabilize one release baseline", "Freeze a named release candidate from dev-main, merge approved branch work, and separate unfinished local changes into module-scoped PRs."),
        ("2", "Prove production readiness", "Apply and register migrations, run role/farm/browser smoke tests, reconcile ledger totals, and capture deployment evidence per module."),
        ("3", "Finish the connected Broiler lifecycle", "Validate DOC Placement → Flock Card → Harvest & Delivery → Clean up → Cycle Master with real farm, building, batch, and warehouse data."),
        ("4", "Operationalize the notification platform", "Complete farm identity, provider configuration, event activation, permission checks, safe no-op, retry, and deduplication verification."),
        ("5", "Institutionalize module ownership", "Maintain the FMS/module/dependency catalog used in this report and assign a product owner, technical owner, deployment status, and next milestone to each module."),
        ("6", "Measure adoption and control quality", "Define CEO-level indicators: active users by FMS, posted documents by module, cycle completion, inventory reconciliation exceptions, approval turnaround, and notification delivery health."),
    ]
    for number, title, body in priorities:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        set_cell_width(table.cell(0, 0), 0.48)
        set_cell_width(table.cell(0, 1), 5.95)
        set_cell_shading(table.cell(0, 0), GOLD)
        set_cell_shading(table.cell(0, 1), LIGHT_GOLD if int(number) % 2 == 0 else LIGHT)
        set_cell_margins(table.cell(0, 0), 100, 80, 100, 80)
        set_cell_margins(table.cell(0, 1), 105, 130, 105, 130)
        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, number, size=15, bold=True, color=WHITE, font="Bahnschrift")
        p = table.cell(0, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_text(p, title, size=10.5, bold=True, color=GREEN, font="Bahnschrift")
        p2 = table.cell(0, 1).add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        add_text(p2, body, size=9.2, color=INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_heading(doc, "Closing perspective", kicker="CEO takeaway")
    add_callout(
        doc,
        "The platform is becoming connected; the next milestone is becoming controlled and provably live.",
        "The codebase now shows the architecture of a unified farm-management platform: common master data, FMS-specific workflows, a shared inventory ledger, farm-scoped access, approvals, notifications, traceability, and management reporting. The highest-return next step is not another broad feature wave—it is disciplined integration, deployment evidence, end-to-end verification, and adoption measurement on the connected workflows already built.",
        fill=LIGHT,
        accent=GREEN,
    )

    add_heading(doc, "Appendix: module naming and evidence notes", kicker="Reference")
    notes = [
        "Vita FMS categories used in this report: Shared Platform & Governance, Inventory, Broiler, Hatchery, and Breeder.",
        "Broiler DOC Placement is implemented under /inv/doc-receiving; its settings are under /a_dean/doc-receiving-settings.",
        "Hatchery Receiving is a separate module under /a_dean/receiving.",
        "Growing & Farm Condition is the current business label for the Broiler Flock Card module under /brd/fc.",
        "Harvest & Delivery is the current business label for BR Delivery under /brd/dr.",
        "Item Stock In and Item Stock Out are the current navigation labels for Goods Receipt (/inv/gr) and Goods Issue (/inv/gi).",
        "Integrated means committed to origin/dev-main as of the refreshed GitHub snapshot; it does not imply production deployment.",
        "Generated poultry-operation photographs in this report are contextual presentation visuals, not evidence of a specific Vitarich site or deployed system state.",
    ]
    for note in notes:
        add_bullet(doc, note)

    add_header_footer(doc)
    doc.core_properties.title = "Vita FMS Total Project Update — May to August 2026"
    doc.core_properties.subject = "CEO project update"
    doc.core_properties.author = "Vita FMS Project Team"
    doc.core_properties.keywords = "Vita FMS, CEO update, Broiler, Breeder, Hatchery, Inventory"
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
